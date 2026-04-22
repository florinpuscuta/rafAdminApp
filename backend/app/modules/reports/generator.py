"""
Generator de raport Word (.docx) pentru dashboard overview.
Structură: header, KPIs, top stores/agents/products, tabel lunar PY vs CY.
"""
from decimal import Decimal
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor


MONTHS_RO = ["Ian","Feb","Mar","Apr","Mai","Iun","Iul","Aug","Sep","Oct","Noi","Dec"]


def _fmt_ron(value: str | float | Decimal) -> str:
    try:
        n = float(value)
    except (TypeError, ValueError):
        return str(value)
    return f"{n:,.2f}".replace(",", " ").replace(".", ",") + " RON"


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    if level == 0:
        for run in h.runs:
            run.font.size = Pt(18)


def _add_kpi_row(doc: Document, kpis: dict, compare_kpis: dict | None, compare_year: int | None) -> None:
    rows = [
        ("Total valoare", _fmt_ron(kpis["total_amount"])),
        ("Linii import", str(kpis["total_rows"])),
        ("Magazine canonice", str(kpis["distinct_mapped_stores"])),
        ("Agenți canonici", str(kpis["distinct_mapped_agents"])),
    ]
    table = doc.add_table(rows=2, cols=len(rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows):
        c_top = table.cell(0, i)
        c_top.text = label
        for run in c_top.paragraphs[0].runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        c_bot = table.cell(1, i)
        c_bot.text = value
        for run in c_bot.paragraphs[0].runs:
            run.font.size = Pt(13)
            run.bold = True
    if compare_kpis and compare_year:
        doc.add_paragraph(
            f"Comparativ cu {compare_year}: total {_fmt_ron(compare_kpis['total_amount'])}, "
            f"{compare_kpis['total_rows']} linii."
        )


def _add_top_table(doc: Document, title: str, rows: list[tuple[str, str, str]]) -> None:
    """rows: list of (primary, secondary, amount_str)."""
    _add_heading(doc, title, level=2)
    if not rows:
        doc.add_paragraph("—")
        return
    table = doc.add_table(rows=len(rows) + 1, cols=3)
    table.style = "Light List Accent 1"
    table.cell(0, 0).text = "Nume"
    table.cell(0, 1).text = "Detalii"
    table.cell(0, 2).text = "Total"
    for i, (primary, secondary, amount) in enumerate(rows, start=1):
        table.cell(i, 0).text = primary
        table.cell(i, 1).text = secondary
        table.cell(i, 2).text = amount


def _add_monthly_table(doc: Document, cy_year: int, py_year: int | None,
                       monthly: list, monthly_compare: list) -> None:
    _add_heading(doc, "Vânzări lunare", level=2)
    cy_map = {m["month"]: m["total_amount"] for m in monthly}
    py_map = {m["month"]: m["total_amount"] for m in monthly_compare} if monthly_compare else {}
    headers = ["Luna", str(cy_year)]
    if py_year is not None:
        headers += [str(py_year), "Δ %"]
    table = doc.add_table(rows=13, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for idx, name in enumerate(MONTHS_RO):
        row = idx + 1
        cy = float(cy_map.get(idx + 1, 0) or 0)
        py = float(py_map.get(idx + 1, 0) or 0)
        table.cell(row, 0).text = name
        table.cell(row, 1).text = _fmt_ron(cy) if cy else "—"
        if py_year is not None:
            table.cell(row, 2).text = _fmt_ron(py) if py else "—"
            delta = ((cy - py) / py * 100) if py else None
            table.cell(row, 3).text = (f"{delta:+.1f}%" if delta is not None else "—")


def generate_dashboard_report(overview: dict, tenant_name: str) -> bytes:
    doc = Document()

    _add_heading(doc, f"Raport vânzări — {tenant_name}", level=0)
    year = overview.get("year")
    month = overview.get("month")
    chain = overview.get("chain")
    scope = overview.get("scope")

    subtitle_parts: list[str] = []
    if year:
        subtitle_parts.append(f"An: {year}")
    if month:
        subtitle_parts.append(f"Luna: {MONTHS_RO[month - 1]}")
    if chain:
        subtitle_parts.append(f"Lanț: {chain}")
    if scope and scope.get("store_name"):
        subtitle_parts.append(f"Magazin: {scope['store_name']}")
    if scope and scope.get("agent_name"):
        subtitle_parts.append(f"Agent: {scope['agent_name']}")
    if scope and scope.get("product_name"):
        subtitle_parts.append(f"Produs: {scope['product_name']}")
    if subtitle_parts:
        p = doc.add_paragraph(" · ".join(subtitle_parts))
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    _add_heading(doc, "KPI", level=2)
    _add_kpi_row(doc, overview["kpis"], overview.get("compare_kpis"), overview.get("compare_year"))

    # Monthly
    if year and overview.get("monthly_totals"):
        _add_monthly_table(
            doc,
            year,
            overview.get("compare_year"),
            overview["monthly_totals"],
            overview.get("monthly_totals_compare") or [],
        )

    # Top chains
    if overview.get("top_chains"):
        _add_top_table(
            doc,
            "Top lanțuri",
            [
                (c["chain"], f"{c['store_count']} magazine", _fmt_ron(c["total_amount"]))
                for c in overview["top_chains"]
            ],
        )

    # Top stores/agents/products
    if overview.get("top_stores"):
        _add_top_table(
            doc,
            "Top magazine",
            [
                (s["store_name"], s["chain"] or "—", _fmt_ron(s["total_amount"]))
                for s in overview["top_stores"]
            ],
        )
    if overview.get("top_agents"):
        _add_top_table(
            doc,
            "Top agenți",
            [
                (a["agent_name"], "", _fmt_ron(a["total_amount"]))
                for a in overview["top_agents"]
            ],
        )
    if overview.get("top_products"):
        _add_top_table(
            doc,
            "Top produse",
            [
                (p["product_name"], p["product_code"] or "—", _fmt_ron(p["total_amount"]))
                for p in overview["top_products"]
            ],
        )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
