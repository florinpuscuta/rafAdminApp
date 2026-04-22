"""Asamblare raport docx complet — structura RAF Consulting (~15-20 pagini).

Secțiuni:
  0. Cover page
  1. Tablou de bord grafic (2 figuri overview)
  2. I. Adeplast KA — Sumar, clienți, grupe, marca privată
  3. II. Sika KA — Sumar, clienți, grupe
  4. III. Analiza consolidată
  5. IV. Concluzii și recomandări strategice
  6. V. Vânzări pe zone geografice
  7. VI. Evoluție lunară (YTD)
  8. Footer
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt, RGBColor

from app.modules.monthly_report import chart as chart_mod

if TYPE_CHECKING:
    from app.modules.monthly_report.service import (
        BrandDossier,
        FullDossier,
        MarcaPrivataDossier,
        Row,
        YoY,
        ZoneDossier,
    )

MONTHS_RO = [
    "", "Ianuarie", "Februarie", "Martie", "Aprilie", "Mai", "Iunie",
    "Iulie", "August", "Septembrie", "Octombrie", "Noiembrie", "Decembrie",
]

COLOR_ADP = "#2563eb"   # albastru
COLOR_SIKA = "#f59e0b"  # portocaliu (galben sika)
COLOR_PL = "#8b5cf6"    # mov (marca privată)
COLOR_TOTAL = "#0ea5e9"


def _fmt_ron(v: float) -> str:
    return f"{v:,.0f}".replace(",", ".")


def _fmt_pct(p: float | None) -> str:
    if p is None:
        return "—"
    sign = "+" if p > 0 else ""
    return f"{sign}{p:.1f}%"


def _pct_color(p: float | None) -> RGBColor:
    if p is None or p == 0:
        return RGBColor(0x64, 0x74, 0x8B)
    return RGBColor(0x05, 0x96, 0x69) if p > 0 else RGBColor(0xDC, 0x26, 0x26)


def _hex_to_rgb(h: str) -> RGBColor:
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _paragraph(doc: Document, text: str, *, size: int = 11,
               bold: bool = False, italic: bool = False,
               color: RGBColor | None = None, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color


def _heading(doc: Document, text: str, *, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _image(doc: Document, png: bytes, *, width_cm: float = 16) -> None:
    try:
        doc.add_picture(io.BytesIO(png), width=Cm(width_cm))
    except Exception:
        _paragraph(doc, "(chart indisponibil)", italic=True, size=10)


def _page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _start_landscape(doc: Document) -> None:
    """Începe o secțiune nouă în landscape A4."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)


def _start_portrait(doc: Document) -> None:
    """Revine la portrait A4 după o secțiune landscape."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)


def _yoy_table(doc: Document, *, header: list[str], rows: list[tuple[str, float, float, float, float | None]]) -> None:
    """Tabel cu 5 coloane: label | cur | prev | diff | pct."""
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for label, cur, prev, diff, pct in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = _fmt_ron(cur)
        row[2].text = _fmt_ron(prev)
        row[3].text = _fmt_ron(diff)
        row[4].text = _fmt_pct(pct)
        for r in row[4].paragraphs[0].runs:
            r.font.color.rgb = _pct_color(pct)
            r.bold = True


# ─────────────────────────────────────────────────────────────────────────
# Builders per section
# ─────────────────────────────────────────────────────────────────────────

def _build_cover(doc: Document, d: "FullDossier") -> None:
    _paragraph(doc, "RAF CONSULTING", size=22, bold=True, color=_hex_to_rgb("#0f172a"),
               align=WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc, "Florin Puscuta — General Manager", size=12, italic=True,
               color=_hex_to_rgb("#64748b"), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().add_run().add_break()
    doc.add_paragraph().add_run().add_break()
    _paragraph(doc, "RAPORT KA MANAGEMENT", size=28, bold=True,
               color=_hex_to_rgb("#2563eb"), align=WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc, "Adeplast & Sika — Canal Key Accounts", size=14, italic=True,
               color=_hex_to_rgb("#334155"), align=WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc, f"{MONTHS_RO[d.month]} {d.year}  |  Analiza Lunară și YTD",
               size=13, color=_hex_to_rgb("#475569"), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().add_run().add_break()
    doc.add_paragraph().add_run().add_break()
    _paragraph(doc, "Confidențial — RAF Consulting", size=10, italic=True,
               color=_hex_to_rgb("#94a3b8"), align=WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc, datetime.now().strftime("%d %B %Y"), size=10,
               color=_hex_to_rgb("#94a3b8"), align=WD_ALIGN_PARAGRAPH.CENTER)


def _build_overview(doc: Document, d: "FullDossier") -> None:
    _heading(doc, "TABLOU DE BORD GRAFIC — Imagine de Ansamblu", level=1)

    # Fig. 1 — comparație branduri
    _paragraph(doc,
        f"Fig. 1 — Comparație Vânzări KA: {MONTHS_RO[d.month]} {d.prev_year} vs {MONTHS_RO[d.month]} {d.year}",
        size=11, bold=True, color=_hex_to_rgb("#475569"))
    png = chart_mod.comparison_bars(
        title=f"Adeplast vs Sika — {MONTHS_RO[d.month]} {d.year} vs {d.prev_year}",
        series=[
            ("Adeplast", d.adeplast.amount.current, d.adeplast.amount.prev, COLOR_ADP),
            ("Sika", d.sika.amount.current, d.sika.amount.prev, COLOR_SIKA),
            ("Total KA", d.total.current, d.total.prev, COLOR_TOTAL),
        ],
        label_current=str(d.year),
        label_prev=str(d.prev_year),
    )
    _image(doc, png)
    doc.add_paragraph()

    # Fig. 2 — structură portofoliu (donut)
    _paragraph(doc, "Fig. 2 — Structura Portofoliu KA",
               size=11, bold=True, color=_hex_to_rgb("#475569"))
    png2 = chart_mod.donut_chart(
        title=f"Pondere pe brand — {MONTHS_RO[d.month]} {d.year}",
        segments=[
            ("Adeplast", d.adeplast.amount.current, COLOR_ADP),
            ("Sika", d.sika.amount.current, COLOR_SIKA),
        ],
    )
    _image(doc, png2, width_cm=14)


def _build_executive_summary(doc: Document, text: str) -> None:
    _heading(doc, "SUMAR EXECUTIV", level=2)
    for para in (text or "(fără narațiune)").split("\n\n"):
        _paragraph(doc, para.strip(), size=11)


def _build_brand_section(
    doc: Document,
    bd: "BrandDossier",
    *,
    brand_color: str,
    narrative: str,
    clients_narrative: str,
    categories_narrative: str,
) -> None:
    title = f"{'I' if bd.brand == 'Adeplast' else 'II'}. {bd.brand.upper()} KEY ACCOUNTS — "
    title += f"{MONTHS_RO[bd.month]} {bd.year}"
    _heading(doc, title, level=1)

    # Sumar executiv brand
    _heading(doc, "Sumar Executiv", level=3)
    for para in (narrative or "").split("\n\n"):
        _paragraph(doc, para.strip(), size=11)

    # KPI row (tabel 2x4)
    kpi_table = doc.add_table(rows=2, cols=4)
    kpi_table.style = "Light Shading Accent 1"
    headers = [
        (f"Luna {MONTHS_RO[bd.month]}", _fmt_ron(bd.amount.current) + " RON"),
        (f"vs {bd.prev_year}", _fmt_pct(bd.amount.pct)),
        (f"YTD {bd.year}", _fmt_ron(bd.amount_ytd.current) + " RON"),
        (f"vs YTD {bd.prev_year}", _fmt_pct(bd.amount_ytd.pct)),
    ]
    for i, (lbl, val) in enumerate(headers):
        kpi_table.rows[0].cells[i].text = lbl
        kpi_table.rows[1].cells[i].text = val
        for r in kpi_table.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(9)
        for r in kpi_table.rows[1].cells[i].paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(11)

    doc.add_paragraph()

    # Top clienți
    _heading(doc, "Vânzări pe Clienți KA", level=3)
    if bd.top_clients:
        for para in (clients_narrative or "").split("\n\n"):
            _paragraph(doc, para.strip(), size=11)
        png = chart_mod.horizontal_bar_chart(
            [chart_mod.BarRow(label=c.label, current=c.cur, prev=c.prev) for c in bd.top_clients[:8]],
            label_current=str(bd.year), label_prev=str(bd.prev_year),
            title=f"Top {len(bd.top_clients[:8])} clienți {bd.brand}",
        )
        _image(doc, png)
        _yoy_table(
            doc,
            header=["Client (Lanț)", str(bd.year), str(bd.prev_year), "Diferență", "%"],
            rows=[(c.label, c.cur, c.prev, c.diff, c.pct) for c in bd.top_clients],
        )
    else:
        _paragraph(doc, "(fără date pentru clienți)", italic=True, size=10)

    doc.add_paragraph()

    # Grupe produse
    _heading(doc, "Grupe de Produse — Top Categorii", level=3)
    if bd.top_categories:
        for para in (categories_narrative or "").split("\n\n"):
            _paragraph(doc, para.strip(), size=11)
        total_cat = sum(c.cur for c in bd.top_categories) or 1.0
        donut_segs = []
        palette = ["#2563eb", "#f59e0b", "#22c55e", "#ef4444", "#8b5cf6",
                   "#06b6d4", "#ec4899", "#84cc16", "#64748b", "#a855f7"]
        for idx, c in enumerate(bd.top_categories[:10]):
            donut_segs.append((c.label, c.cur, palette[idx % len(palette)]))
        png = chart_mod.donut_chart(
            title=f"Structura vânzări pe grupe {bd.brand}",
            segments=donut_segs,
        )
        _image(doc, png, width_cm=15)
        _yoy_table(
            doc,
            header=["Categorie", str(bd.year), str(bd.prev_year), "Diferență", "%"],
            rows=[(c.label, c.cur, c.prev, c.diff, c.pct) for c in bd.top_categories],
        )
    else:
        _paragraph(doc, "(fără date pentru categorii)", italic=True, size=10)

    # marker color (not strictly needed; keep placeholder for future use)
    _ = brand_color


def _build_marca_privata(
    doc: Document, mp: "MarcaPrivataDossier", *, narrative: str,
) -> None:
    _heading(doc, "Marcă Privată — Performanță", level=2)
    for para in (narrative or "").split("\n\n"):
        _paragraph(doc, para.strip(), size=11)

    # Chart donut ADP vs PL
    png = chart_mod.donut_chart(
        title=f"Adeplast brand vs Marcă Privată — {MONTHS_RO[mp.month]} {mp.year}",
        segments=[
            ("Adeplast", mp.adeplast.current, COLOR_ADP),
            ("Marcă Privată", mp.marca_privata.current, COLOR_PL),
        ],
    )
    _image(doc, png, width_cm=14)

    _yoy_table(
        doc,
        header=["", str(mp.year), str(mp.prev_year), "Diferență", "%"],
        rows=[
            ("Adeplast brand", mp.adeplast.current, mp.adeplast.prev,
             mp.adeplast.current - mp.adeplast.prev, mp.adeplast.pct),
            ("Marcă Privată", mp.marca_privata.current, mp.marca_privata.prev,
             mp.marca_privata.current - mp.marca_privata.prev, mp.marca_privata.pct),
        ],
    )

    if mp.categories_pl:
        _paragraph(doc, "Categorii Marcă Privată — detaliu:", bold=True, size=11)
        _yoy_table(
            doc,
            header=["Categorie PL", str(mp.year), str(mp.prev_year), "Diferență", "%"],
            rows=[(c.label, c.cur, c.prev, c.diff, c.pct) for c in mp.categories_pl],
        )


def _build_consolidated(
    doc: Document, d: "FullDossier", *, narrative: str,
) -> None:
    _heading(doc, "III. ANALIZA CONSOLIDATĂ — Adeplast + Sika KA", level=1)

    for para in (narrative or "").split("\n\n"):
        _paragraph(doc, para.strip(), size=11)

    # Top clienți consolidați
    if d.consolidated_top_clients:
        _heading(doc, "Top Clienți Consolidați", level=3)
        png = chart_mod.horizontal_bar_chart(
            [chart_mod.BarRow(label=c.label, current=c.cur, prev=c.prev)
             for c in d.consolidated_top_clients[:10]],
            label_current=str(d.year), label_prev=str(d.prev_year),
            title=f"Top {len(d.consolidated_top_clients[:10])} clienți KA consolidați — {MONTHS_RO[d.month]} {d.year}",
        )
        _image(doc, png)
        _yoy_table(
            doc,
            header=["Client (Lanț)", str(d.year), str(d.prev_year), "Diferență", "%"],
            rows=[(c.label, c.cur, c.prev, c.diff, c.pct) for c in d.consolidated_top_clients],
        )


def _build_conclusions(doc: Document, *, narrative: str) -> None:
    _heading(doc, "IV. CONCLUZII ȘI RECOMANDĂRI STRATEGICE", level=1)
    for para in (narrative or "").split("\n\n"):
        _paragraph(doc, para.strip(), size=11)


def _build_zones(doc: Document, z: "ZoneDossier", *, narrative: str, top_n: int = 15) -> None:
    _heading(doc, "V. VÂNZĂRI PE ZONE GEOGRAFICE", level=1)
    for para in (narrative or "").split("\n\n"):
        _paragraph(doc, para.strip(), size=11)

    top = z.zones[:top_n]
    if top:
        png = chart_mod.horizontal_bar_chart(
            [chart_mod.BarRow(label=r.zone, current=r.amount_current, prev=r.amount_prev)
             for r in top],
            label_current=str(z.year), label_prev=str(z.prev_year),
            title=f"Top {len(top)} zone — {MONTHS_RO[z.month]} {z.year} vs {z.prev_year}",
        )
        _image(doc, png)
        _yoy_table(
            doc,
            header=["Zonă", str(z.year), str(z.prev_year), "Diferență", "%"],
            rows=[(r.zone, r.amount_current, r.amount_prev, r.diff, r.pct) for r in top],
        )
        # Total
        tp = doc.add_paragraph()
        tp.add_run(f"TOTAL (toate zonele): {_fmt_ron(z.total_current)} RON  ").bold = True
        run = tp.add_run(_fmt_pct(z.total_pct))
        run.bold = True
        run.font.color.rgb = _pct_color(z.total_pct)


def _build_monthly_evolution(doc: Document, d: "FullDossier") -> None:
    _heading(doc, "VI. EVOLUȚIA LUNARĂ YTD", level=1)
    _paragraph(doc,
        f"Evoluția vânzărilor lunare pentru perioada Ianuarie – "
        f"{MONTHS_RO[d.month]} {d.year} comparativ cu {d.prev_year}.",
        size=11,
    )
    months = [m for m, _, _ in d.monthly_evolution]
    cur = [c for _, c, _ in d.monthly_evolution]
    prev = [p for _, _, p in d.monthly_evolution]
    png = chart_mod.evolution_lines(
        title=f"Evoluție lunară KA — {d.year} vs {d.prev_year}",
        months=months,
        series=[
            (str(d.year), cur, COLOR_TOTAL),
            (str(d.prev_year), prev, "#94a3b8"),
        ],
    )
    _image(doc, png)


def _build_top_products(
    doc: Document, bd: "BrandDossier", *, heading_num: str,
) -> None:
    _heading(doc, f"{heading_num} TOP 15 PRODUSE — {bd.brand.upper()}", level=1)
    if not bd.top_products:
        _paragraph(doc, "(fără produse cu vânzări în luna aleasă)",
                   italic=True, size=10)
        return

    _paragraph(doc,
        f"Clasamentul primelor 15 produse {bd.brand} din canalul KA, "
        f"cu comparație {bd.year} vs {bd.prev_year}.", size=11,
    )

    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr_labels = ["#", "Cod", "Produs", "Vânzări " + str(bd.year),
                  str(bd.prev_year), "%"]
    for i, h in enumerate(hdr_labels):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True; r.font.size = Pt(10)

    for idx, pr in enumerate(bd.top_products, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = (pr.code or "")[:14]
        row[2].text = (pr.name or "")[:50]
        row[3].text = _fmt_ron(pr.cur)
        row[4].text = _fmt_ron(pr.prev)
        row[5].text = _fmt_pct(pr.pct)
        for r in row[5].paragraphs[0].runs:
            r.font.color.rgb = _pct_color(pr.pct); r.bold = True

    # Chart top 15 under table (table first, per user request)
    doc.add_paragraph()
    png = chart_mod.horizontal_bar_chart(
        [chart_mod.BarRow(label=(pr.name[:22] if pr.name else pr.code),
                          current=pr.cur, prev=pr.prev)
         for pr in bd.top_products],
        label_current=str(bd.year), label_prev=str(bd.prev_year),
        title=f"Top 15 produse {bd.brand} — {bd.year} vs {bd.prev_year}",
    )
    _image(doc, png)


def _build_price_tables_company(
    doc: Document, *, anchor_label: str, tables, chart_png: bytes | None,
) -> None:
    """Un subsection per magazin cu tabel side-by-side + legendă + chart la final."""
    if not tables:
        _paragraph(doc, f"(fără date de preț pentru {anchor_label})",
                   italic=True, size=10)
        return

    for tbl in tables:
        _paragraph(doc,
            f"{anchor_label} — {tbl.store}  ({len(tbl.rows)} produse comparate)",
            size=13, bold=True, color=_hex_to_rgb("#334155"),
        )
        # Simplificat: 1 coloană (produs + preț concatenat) per brand, nu 2
        ncols = 1 + len(tbl.columns)
        t = doc.add_table(rows=1, cols=ncols)
        t.style = "Light Grid Accent 1"
        t.autofit = True
        hdr = t.rows[0].cells
        hdr[0].text = "#"
        for i, brand in enumerate(tbl.columns):
            hdr[i + 1].text = brand.upper() if brand.lower() == "adeplast" else brand
        for c in hdr:
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True; r.font.size = Pt(9)

        for pline in tbl.rows:
            row = t.add_row().cells
            row[0].text = str(pline.row_num)
            for p in row[0].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
            for i, brand in enumerate(tbl.columns):
                obj = pline.brand_prices.get(brand, {})
                prod = (obj.get("prod") or "") if isinstance(obj, dict) else ""
                pret = obj.get("pret") if isinstance(obj, dict) else None
                cell = row[i + 1]
                cell.text = ""
                para = cell.paragraphs[0]
                # Linia 1: produs (font mic)
                r1 = para.add_run(str(prod)[:28] if prod else "—")
                r1.font.size = Pt(7.5)
                r1.font.color.rgb = _hex_to_rgb("#475569")
                # Linia 2: preț (font mare, accent)
                para.add_run("\n")
                r2 = para.add_run(
                    _fmt_ron(float(pret)) + " RON"
                    if isinstance(pret, (int, float)) and pret > 0 else "—"
                )
                r2.font.size = Pt(10)
                if brand == tbl.anchor_brand:
                    r2.bold = True
                    r2.font.color.rgb = _hex_to_rgb("#2563eb")
                else:
                    r2.font.color.rgb = _hex_to_rgb("#0f172a")
        doc.add_paragraph()

    if chart_png is not None:
        _paragraph(doc, f"Grafic de ansamblu — {anchor_label}", size=11,
                   bold=True, color=_hex_to_rgb("#475569"))
        _image(doc, chart_png)


def _build_prices(doc: Document, pd, *, narrative: str) -> None:
    _heading(doc, "VII. ANALIZA PREȚURI — Adeplast & Sika vs Competitori", level=1)
    for para in (narrative or "").split("\n\n"):
        _paragraph(doc, para.strip(), size=11)

    if pd.avg_advantage_pct is not None:
        p = doc.add_paragraph()
        p.add_run("Avantaj mediu Adeplast vs competitori direcți: ").bold = True
        color = _pct_color(-pd.avg_advantage_pct)
        r = p.add_run(_fmt_pct(pd.avg_advantage_pct))
        r.bold = True
        r.font.color.rgb = color
    if pd.stores_covered:
        _paragraph(doc,
            f"Retaileri cuprinși: {', '.join(pd.stores_covered)}.",
            size=11,
        )
    doc.add_paragraph()

    # ─── VII.1 Adeplast în toate lanțurile — tabele side-by-side ──────────
    _heading(doc, "VII.1 ADEPLAST în toate lanțurile KA", level=2)
    _paragraph(doc,
        "Tabelele de mai jos arată produsele Adeplast (coloană evidențiată) "
        "alături de concurenții direcți pe aceeași linie (Sika, Mapei, Baumit, "
        "Ceresit, 4 Maini, mărci proprii). Analiza se face pe LINIE — fiecare "
        "rând e un produs comparat cu concurenții săi direcți în același magazin.",
        size=11,
    )
    doc.add_paragraph()
    # Chart sumar Adeplast după tabele
    adp_chart = None
    if pd.rows:
        adp_chart = chart_mod.horizontal_bar_chart(
            [chart_mod.BarRow(label=f"{pr.store}: {pr.product_adp[:16]}",
                              current=pr.price_adp,
                              prev=pr.avg_competitor or 0)
             for pr in pd.rows[:12]],
            label_current="Adeplast", label_prev="Media competitori",
            title="Adeplast vs media competitori — top 12 linii",
        )
    _build_price_tables_company(
        doc, anchor_label="Adeplast",
        tables=pd.adeplast_tables, chart_png=adp_chart,
    )

    # ─── VII.2 Sika în toate lanțurile ────────────────────────────────────
    _heading(doc, "VII.2 SIKA în toate lanțurile KA", level=2)
    _paragraph(doc,
        "Tabelele de mai jos arată produsele Sika (coloană evidențiată) "
        "alături de concurenții direcți pe aceeași linie (Bostik, Soudal, "
        "Ceresit, Weber, Mapei, Adeplast). Analiza se face pe linie — "
        "produse comparative în același magazin.",
        size=11,
    )
    doc.add_paragraph()
    _build_price_tables_company(
        doc, anchor_label="Sika",
        tables=pd.sika_tables, chart_png=None,
    )


def _build_marketing(doc: Document, md, *, narrative: str) -> None:
    _heading(doc, "VIII. ACTIVITĂȚI MARKETING", level=1)
    for para in (narrative or "").split("\n\n"):
        _paragraph(doc, para.strip(), size=11)

    # Sumar volume
    _paragraph(doc,
        f"Inventar curent: {md.panouri_count} poze panouri · "
        f"{md.magazine_count} poze magazine/concurență · "
        f"{md.catalog_count} poze catalog.",
        size=11, italic=True,
    )

    # Catalog Lunar
    if md.catalog_photos:
        _heading(doc,
            f"VIII.1 Catalog Lunar — {md.catalog_folder_name or 'luna curentă'}",
            level=3,
        )
        for photo in md.catalog_photos:
            try:
                doc.add_picture(io.BytesIO(photo.png_bytes), width=Cm(12))
                _paragraph(doc, photo.caption, size=9, italic=True,
                           color=_hex_to_rgb("#64748b"))
            except Exception:
                continue

    # Panouri
    if md.panouri_photos:
        _heading(doc, "VIII.2 Panouri & Standuri — Exemple Active", level=3)
        for photo in md.panouri_photos:
            try:
                doc.add_picture(io.BytesIO(photo.png_bytes), width=Cm(12))
                _paragraph(doc, photo.caption, size=9, italic=True,
                           color=_hex_to_rgb("#64748b"))
            except Exception:
                continue

    # Magazine / Concurență
    if md.magazine_photos:
        _heading(doc, "VIII.3 Prezență la Raft și Concurență", level=3)
        for photo in md.magazine_photos:
            try:
                doc.add_picture(io.BytesIO(photo.png_bytes), width=Cm(12))
                _paragraph(doc, photo.caption, size=9, italic=True,
                           color=_hex_to_rgb("#64748b"))
            except Exception:
                continue


def _bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(line, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(11)


def _sep(doc: Document) -> None:
    """Linie separatoare (hairline)."""
    p = doc.add_paragraph()
    r = p.add_run("─" * 60)
    r.font.color.rgb = _hex_to_rgb("#cbd5e1")
    r.font.size = Pt(10)


def _build_activities(doc: Document, d: "FullDossier") -> None:
    """Secțiune fixă — activități recurente (text static, cu luna substituită)."""
    month_name = MONTHS_RO[d.month].lower()
    year = d.year

    _heading(doc, f"8. ACTIVITĂȚI DESFĂȘURATE — ADEPLAST", level=1)
    _sep(doc)

    _heading(doc, f"8.1 Activități recurente desfășurate în luna {month_name} {year}", level=3)
    _paragraph(doc,
        "Alte acțiuni realizate în această lună sunt de fapt recurente și se "
        "desfășoară lună de lună, fără întrerupere, constituind fundamentul "
        "execuției comerciale în teritoriu.",
        size=11,
    )
    _paragraph(doc,
        "Aceste activități sunt esențiale și reprezintă diferența dintre o lună "
        "medie și o lună pozitivă în condiții dificile.",
        size=11,
    )

    # Analiza sortimentului
    _paragraph(doc, "Analiza sortimentului la fiecare punct de vânzare", bold=True, size=11)
    _paragraph(doc,
        "În fiecare zi de luni a săptămânii se trasează obiectivele specifice "
        "fiecărui magazin / zonă / merchandiser / agent.",
        size=11,
    )
    _paragraph(doc, "Aceasta presupune:", size=11)
    _bullets(doc, [
        "analiză stoc vs rulaj",
        "identificare produse lente",
        "identificare lipsuri în raft",
        "verificare corelare promo vs stoc",
    ])
    _paragraph(doc,
        "Această activitate se desfășoară săptămânal și este monitorizată direct.",
        size=11,
    )
    _sep(doc)

    # Analiza profitabilității
    _paragraph(doc, "Analiza profitabilității fiecărui produs", bold=True, size=11)
    _paragraph(doc,
        "Se verifică marjele reale obținute pe fiecare SKU comercializat.",
        size=11,
    )
    _paragraph(doc, "Se urmărește:", size=11)
    _bullets(doc, [
        "impact discounturi",
        "impact cost logistic",
        "impact rotație lentă",
        "identificare produse neprofitabile",
    ])
    _sep(doc)

    # Analiza concurenței
    _paragraph(doc, "Analiza concurenței", bold=True, size=11)
    _paragraph(doc,
        "Această activitate se desfășoară continuu și presupune:",
        size=11,
    )
    _bullets(doc, [
        "monitorizare prețuri concurență (Ceresit, Baumit, Mapei etc.)",
        "verificare fețe de raft ocupate",
        "identificare acțiuni promo active",
        "monitorizare implementări noi în magazine",
    ])
    _sep(doc)

    # Analiza acțiunilor Adeplast
    _paragraph(doc, "Analiza acțiunilor în desfășurare de către ADEPLAST", bold=True, size=11)
    _paragraph(doc,
        "În fiecare zi de marți a săptămânii se face analiză internă:",
        size=11,
    )
    _bullets(doc, [
        "situație stoc",
        "situație livrări",
        "status reclamații",
        "situație implementare materiale promoționale",
    ])
    _sep(doc)

    # Inventariere paleți
    _paragraph(doc, "Inventariere și recuperare paleți", bold=True, size=11)
    _paragraph(doc,
        "Responsabili: agenții de vânzări și asistent manager.",
        size=11,
    )
    _paragraph(doc,
        "Este vorba de fiecare locație din cele 133 puncte de vânzare gestionate.",
        size=11,
    )
    _paragraph(doc,
        "Această activitate este permanentă și esențială pentru controlul "
        "costurilor logistice.",
        size=11,
    )
    _sep(doc)

    # Probleme logistice
    _paragraph(doc, "Analiza și rezolvarea problemelor logistice", bold=True, size=11)
    _paragraph(doc, "Se desfășoară de luni până vineri.", size=11)
    _paragraph(doc, "Include:", size=11)
    _bullets(doc, [
        "verificare comenzi livrate vs comandate",
        "verificare produse deteriorate",
        "gestionare reclamații ambalaj",
        "intervenții rapide la nivel de depozit",
    ])
    _sep(doc)

    # Analiza stocurilor
    _paragraph(doc, "Analiza stocurilor și generarea comenzilor", bold=True, size=11)
    _paragraph(doc,
        "Se desfășoară în colaborare cu șefii de raion ai partenerilor "
        "(luni, miercuri, joi).",
        size=11,
    )
    _paragraph(doc, "Se urmărește:", size=11)
    _bullets(doc, [
        "menținere stoc optim",
        "prevenire rupturi",
        "prevenire suprastoc",
    ])
    _sep(doc)

    # Plasări suplimentare
    _paragraph(doc, "Plasări suplimentare în punctele de vânzare", bold=True, size=11)
    _paragraph(doc,
        "Se realizează în special pentru produsele care vin cu minusuri din "
        "luna anterioară.",
        size=11,
    )
    _paragraph(doc, "Include:", size=11)
    _bullets(doc, [
        "capete de raft",
        "paleți suplimentari",
        "zone secundare de expunere",
    ])
    _sep(doc)

    # Merchandising
    _paragraph(doc, "Acțiuni de merchandising", bold=True, size=11)
    _paragraph(doc, "Conform fișei de post a merchandiserilor.", size=11)
    _bullets(doc, [
        "refacere raft",
        "aliniere etichete",
        "verificare poziționare corectă",
        "eliminare produse expirate",
    ])
    _sep(doc)

    # Asistare DIY
    _paragraph(doc, "Asistarea vânzătorilor DIY la punctul de vânzare", bold=True, size=11)
    _paragraph(doc, "Activitate continuă:", size=11)
    _bullets(doc, [
        "explicații tehnice",
        "suport aplicare",
        "clarificare diferențe produse",
        "susținere recomandare activă",
    ])

    _page_break(doc)

    # ─── Secțiunea 9. SIKA ────────────────────────────────────────────────
    _heading(doc, "9. ACTIVITĂȚI DESFĂȘURATE — SIKA", level=1)
    _sep(doc)

    _heading(doc, f"9.1 Activități recurente desfășurate în luna {month_name} {year}", level=3)
    _paragraph(doc,
        "Structura este identică, cu specific pe gama SIKA.",
        size=11,
    )
    _paragraph(doc,
        "Se repetă aceleași activități operaționale, adaptate la:",
        size=11,
    )
    _bullets(doc, [
        "segmentele TM",
        "poziționarea diferită în raft",
        "focus pe sealing & bonding",
    ])
    _sep(doc)

    _paragraph(doc, "Analiza sortimentului SIKA", bold=True, size=11)
    _paragraph(doc, "Se verifică în special:", size=11)
    _bullets(doc, [
        "poziționarea produselor 107 / 255 / 11FC",
        "vizibilitate hidro",
        "rotație produse tehnice",
    ])
    _sep(doc)

    _paragraph(doc, "Analiza TM", bold=True, size=11)
    _paragraph(doc, "Se urmărește:", size=11)
    _bullets(doc, [
        "evoluția Building Finishing",
        "scăderea Waterproofing",
        "stagnarea Sealing & Bonding",
    ])
    _sep(doc)

    _paragraph(doc, "Acțiuni în magazine", bold=True, size=11)
    _bullets(doc, [
        "refacere raft chituri și adezivi",
        "verificare materiale promo",
        "susținere produse noi (1045 / 1030)",
    ])
    _sep(doc)

    _paragraph(doc, "Discuții privind participarea în reviste", bold=True, size=11)
    _paragraph(doc,
        "În contextul neimplementării scumpirilor, nu au fost acceptate acțiuni "
        "majore în reviste.",
        size=11,
    )
    _paragraph(doc,
        "Acest lucru a limitat vizibilitatea în luna analizată.",
        size=11,
    )
    _sep(doc)

    _paragraph(doc, "Asistare agenți în ofertare șantiere", bold=True, size=11)
    _paragraph(doc,
        "Segmentul B2B rămâne prioritar pentru susținerea cifrei în perioadele "
        "cu trafic slab DIY.",
        size=11,
    )

    _page_break(doc)

    # ─── Concluzie finală ─────────────────────────────────────────────────
    _heading(doc, "CONCLUZIE FINALĂ — ACTIVITĂȚI", level=1)
    _paragraph(doc,
        "Activitățile recurente reprezintă coloana vertebrală a execuției "
        "comerciale.",
        size=11,
    )
    _paragraph(doc,
        "Într-o lună în care piața nu a ajutat, diferența a fost făcută de:",
        size=11,
    )
    _bullets(doc, [
        "disciplină operațională",
        "prezență în magazin",
        "negociere centrală",
        "menținerea presiunii comerciale",
    ])


def build_screenshots_only_docx(shots, *, title: str = "Capturi Pagini Aplicație") -> bytes:
    """Construiește un docx minimal — doar cover + screenshots (fără date/AI)."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Cover scurt
    _paragraph(doc, "RAF CONSULTING", size=18, bold=True,
               color=_hex_to_rgb("#0f172a"), align=WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc, title, size=22, bold=True,
               color=_hex_to_rgb("#2563eb"), align=WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc,
        f"{datetime.now().strftime('%d %B %Y · %H:%M')}", size=11,
        color=_hex_to_rgb("#64748b"), align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _paragraph(doc,
        f"{len(shots)} pagini capturate live din aplicația Krossdash.",
        size=11, italic=True, color=_hex_to_rgb("#475569"),
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _page_break(doc)

    # Fiecare screenshot pe landscape A4
    _start_landscape(doc)
    for i, shot in enumerate(shots, start=1):
        _paragraph(doc, f"{i}. {shot.label}",
                   size=13, bold=True, color=_hex_to_rgb("#2563eb"))
        _paragraph(doc, f"Path: {shot.path}",
                   size=9, italic=True, color=_hex_to_rgb("#64748b"))
        try:
            doc.add_picture(io.BytesIO(shot.png_bytes), width=Cm(25))
        except Exception as e:
            _paragraph(doc, f"(screenshot invalid: {e})", italic=True, size=10)
        if i < len(shots):
            _page_break(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_app_screenshots(doc: Document, shots) -> None:
    """Anexă — capturi din aplicație (fiecare pe câte o pagină landscape)."""
    if not shots:
        return
    _start_landscape(doc)
    _heading(doc, "ANEXĂ — Capturi din aplicație (date live)", level=1)
    _paragraph(doc,
        f"Mai jos sunt {len(shots)} capturi directe din aplicația Krossdash, "
        "reflectând starea paginilor la momentul generării raportului. "
        "Toate tabelele și graficele sunt extrase 1:1 din UI.",
        size=11,
    )
    _page_break(doc)

    for i, shot in enumerate(shots, start=1):
        _paragraph(doc, f"{i}. {shot.label}",
                   size=13, bold=True, color=_hex_to_rgb("#2563eb"))
        _paragraph(doc, f"Path: {shot.path}",
                   size=9, italic=True, color=_hex_to_rgb("#64748b"))
        try:
            # Width scalat la pagina landscape A4 (~25 cm util).
            # Dacă screenshot-ul e foarte înalt, lăsăm să se spargă natural.
            doc.add_picture(io.BytesIO(shot.png_bytes), width=Cm(24.5))
        except Exception as e:
            _paragraph(doc, f"(screenshot invalid: {e})", italic=True, size=10)
        _page_break(doc)
    _start_portrait(doc)


def _build_footer(doc: Document) -> None:
    doc.add_paragraph()
    _paragraph(doc, "— Sfârșit raport —", italic=True, size=10,
               color=_hex_to_rgb("#94a3b8"), align=WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc,
        f"Generat automat la {datetime.now().strftime('%Y-%m-%d %H:%M')} · "
        "Raport Lunar Management — RAF Consulting",
        italic=True, size=9, color=_hex_to_rgb("#94a3b8"),
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )


# ─────────────────────────────────────────────────────────────────────────
# Public entry point
# ─────────────────────────────────────────────────────────────────────────

def build_full_report_docx(
    d: "FullDossier",
    *,
    narratives: dict[str, str],
    app_screenshots: list | None = None,
) -> bytes:
    """Construiește raportul complet.

    narratives keys:
      exec_summary, adeplast_brand, adeplast_clients, adeplast_categories,
      marca_privata, sika_brand, sika_clients, sika_categories,
      consolidated, conclusions, zones.
    """
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    _build_cover(doc, d)
    _page_break(doc)

    _build_executive_summary(doc, narratives.get("exec_summary", ""))
    _build_overview(doc, d)
    _page_break(doc)

    _build_brand_section(
        doc, d.adeplast,
        brand_color=COLOR_ADP,
        narrative=narratives.get("adeplast_brand", ""),
        clients_narrative=narratives.get("adeplast_clients", ""),
        categories_narrative=narratives.get("adeplast_categories", ""),
    )
    _build_marca_privata(
        doc, d.marca_privata, narrative=narratives.get("marca_privata", ""),
    )
    _page_break(doc)

    _build_brand_section(
        doc, d.sika,
        brand_color=COLOR_SIKA,
        narrative=narratives.get("sika_brand", ""),
        clients_narrative=narratives.get("sika_clients", ""),
        categories_narrative=narratives.get("sika_categories", ""),
    )
    _page_break(doc)

    _build_consolidated(doc, d, narrative=narratives.get("consolidated", ""))
    _page_break(doc)

    _build_conclusions(doc, narrative=narratives.get("conclusions", ""))
    _page_break(doc)

    _build_zones(doc, d.zones, narrative=narratives.get("zones", ""))
    _page_break(doc)

    _build_monthly_evolution(doc, d)

    # ─── Landscape pentru secțiunea Prețuri (tabele late) ────────────────
    _start_landscape(doc)
    _build_prices(doc, d.prices, narrative=narratives.get("prices", ""))
    _start_portrait(doc)

    _build_top_products(doc, d.adeplast, heading_num="VIII.")
    _page_break(doc)
    _build_top_products(doc, d.sika, heading_num="IX.")
    _page_break(doc)

    _build_marketing(doc, d.marketing, narrative=narratives.get("marketing", ""))
    _page_break(doc)

    _build_activities(doc, d)

    if app_screenshots:
        _build_app_screenshots(doc, app_screenshots)

    _build_footer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
